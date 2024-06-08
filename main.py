import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from docxtpl import DocxTemplate
from docx import Document
from docxcompose.composer import Composer
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from tkcalendar import DateEntry
import os
import sys
import logging

# Configure logging
logging.basicConfig(filename='app.log', level=logging.DEBUG,
                    format='%(asctime)s:%(levelname)s:%(message)s')


def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    try:
        base_path = getattr(sys, '_MEIPASS', os.path.abspath("."))
    except AttributeError:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)


# Functions for GUI interactions
def update_pallet_count():
    entry_amount_var.set(str(len(pallet_fields)))


def browse_top_barcode():
    global top_barcode_path
    filepath = filedialog.askopenfilename(initialdir="/", title="Select a File",
                                          filetypes=(("Image files", "*.png *.jpg *.jpeg"), ("all files", "*.*")))
    if filepath:
        top_barcode_path = filepath
        top_barcode_label.config(text=filepath)


def add_pallet_field():
    art = tk.StringVar()
    art_name = tk.StringVar()
    art_amount = tk.StringVar()
    bottom_barcode_var = tk.StringVar()

    pallet_fields.append((art, art_name, art_amount, bottom_barcode_var))
    row = len(pallet_fields) + 8

    tk.Label(root, text="Art").grid(row=row, column=0, pady=2)
    entry_art = tk.Entry(root, textvariable=art)
    entry_art.grid(row=row, column=1, pady=2)

    tk.Label(root, text="Art Name").grid(row=row, column=2, pady=2)
    entry_art_name = tk.Entry(root, textvariable=art_name)
    entry_art_name.grid(row=row, column=3, pady=2)

    tk.Label(root, text="Art Amount").grid(row=row, column=4, pady=2)
    entry_art_amount = tk.Entry(root, textvariable=art_amount)
    entry_art_amount.grid(row=row, column=5, pady=2)

    bottom_barcode_button = tk.Button(root, text="Bottom Barcode",
                                      command=lambda var=bottom_barcode_var: browse_bottom_barcode(var))
    bottom_barcode_button.grid(row=row, column=6, pady=2)

    update_pallet_count()


def browse_bottom_barcode(var):
    filepath = filedialog.askopenfilename(initialdir="/", title="Select a File",
                                          filetypes=(("Image files", "*.png *.jpg *.jpeg"), ("all files", "*.*")))
    if filepath:
        var.set(filepath)


# Function to create a document with custom margins
def create_document_with_custom_margins():
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0)
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1.5)
    return doc


def insert_images_and_text(doc, image_path_top, image_path_bottom, paragraphs, pallet_info):
    # Insert image at the top with adjusted size
    if image_path_top:
        top_run = doc.add_paragraph().add_run()
        top_run.add_picture(image_path_top, width=Inches(5.5), height=Inches(2.3))
        top_run.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Insert the rendered text paragraphs with reduced spacing
    for paragraph in paragraphs:
        new_paragraph = doc.add_paragraph()
        new_paragraph.paragraph_format.space_after = Pt(0)
        new_paragraph.paragraph_format.space_before = Pt(0)
        for run in paragraph.runs:
            new_run = new_paragraph.add_run(run.text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            new_run.font.size = run.font.size
            new_run.font.name = run.font.name
        new_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Insert image at the bottom with adjusted size
    if image_path_bottom:
        bottom_paragraph = doc.add_paragraph()
        bottom_paragraph.paragraph_format.space_before = Pt(0)  # Set space before to 0 to reduce gap
        bottom_paragraph.paragraph_format.space_after = Pt(0)  # Ensure no extra space after
        bottom_run = bottom_paragraph.add_run()
        bottom_run.add_picture(image_path_bottom, width=Inches(5.5), height=Inches(2.3))
        bottom_run.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Insert pallet info
    pallet_paragraph = doc.add_paragraph()
    pallet_run = pallet_paragraph.add_run(pallet_info)
    pallet_run.bold = True
    pallet_run.font.size = Pt(18)
    pallet_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pallet_paragraph.paragraph_format.space_before = Pt(0)  # Ensure no extra space before
    pallet_paragraph.paragraph_format.space_after = Pt(0)  # Ensure no extra space after


def submit_form():
    try:
        global top_barcode_path
        delivery_num = entry_delivery_num.get()
        dest_warehouse = entry_dest_warehouse.get()
        delivery_type = delivery_type_var.get()
        ip = entry_ip.get()
        delivery_date = entry_delivery_date.get()

        # Temporary document files
        temp_docs = []

        for index, (art_var, art_name_var, art_amount_var, bottom_barcode_var) in enumerate(pallet_fields, start=1):
            art = art_var.get()
            art_name = art_name_var.get()
            art_amount = art_amount_var.get()
            bottom_barcode_path = bottom_barcode_var.get()

            context = {
                "delivery_number": delivery_num,
                "ip": ip,
                "warehouse": dest_warehouse,
                "delivery_type": delivery_type,
                "date": delivery_date,
                "pallets": len(pallet_fields),
                "num_of_pall": index,
                "out_of": len(pallet_fields),
                "art": art,
                "art_name": art_name,
                "art_amount": art_amount
            }

            # Load the template and render the context
            doc_tpl = DocxTemplate(resource_path("шаблон3.docx"))
            doc_tpl.render(context)
            temp_path = resource_path(f"temp_rendered_{delivery_num}_{index}.docx")
            doc_tpl.save(temp_path)

            if not os.path.exists(temp_path):
                logging.error(f"Temporary file was not created: {temp_path}")
                messagebox.showerror("Error", f"Temporary file was not created: {temp_path}")
                return

            # Save rendered text to a string, preserving formatting
            rendered_doc = Document(temp_path)
            rendered_paragraphs = rendered_doc.paragraphs
            os.remove(temp_path)

            # Create pallet info string
            pallet_info = f"Паллет {index} из {len(pallet_fields)}"

            # Create a new document and insert images, rendered text, and pallet info
            doc = create_document_with_custom_margins()
            insert_images_and_text(doc, top_barcode_path, bottom_barcode_path, rendered_paragraphs, pallet_info)
            output_temp_path = resource_path(f"temp_Поставка_{delivery_num}_{index}.docx")
            doc.save(output_temp_path)

            if not os.path.exists(output_temp_path):
                logging.error(f"Output temporary file was not created: {output_temp_path}")
                messagebox.showerror("Error", f"Output temporary file was not created: {output_temp_path}")
                return

            temp_docs.append(output_temp_path)

        # Merge documents
        output_path = resource_path(f"Поставка_{delivery_num}.docx")
        merge_documents(output_path, *temp_docs)
        logging.info(f"Merged document has been saved to {output_path}")
        messagebox.showinfo("Success", f"Merged document has been saved to {output_path}")

        # Clean up temporary files
        for temp_doc_path in temp_docs:
            os.remove(temp_doc_path)

    except Exception as e:
        logging.error("Error in submit_form: %s", str(e))
        messagebox.showerror("Error", f"An error occurred: {e}")


# Function to merge documents
def merge_documents(output_path, *docs):
    master_doc = Document(docs[0])
    composer = Composer(master_doc)
    for doc in docs[1:]:
        composer.append(Document(doc))
    composer.save(output_path)


# GUI setup
root = tk.Tk()
root.title("Fill and Download Form")

tk.Label(root, text="Номер поставки").grid(row=0, column=0, padx=10, pady=2)
entry_delivery_num = tk.Entry(root)
entry_delivery_num.grid(row=0, column=1, padx=10, pady=2)

tk.Label(root, text="ИП").grid(row=1, column=0, padx=10, pady=2)
entry_ip = tk.StringVar(root)
entry_ip_options = ttk.Combobox(root, textvariable=entry_ip, state='readonly')
entry_ip_options['values'] = ('Бобоев Э.И.', 'Отабеков Ю.А.', 'Отабеков Ю.М.')
entry_ip_options.grid(row=1, column=1, padx=10, pady=2)

tk.Label(root, text="Склад назначения").grid(row=2, column=0, padx=10, pady=2)
entry_dest_warehouse = tk.StringVar(root)
entry_dest_warehouse_options = ttk.Combobox(root, textvariable=entry_dest_warehouse, state='readonly')
entry_dest_warehouse_options['values'] = (
'Электросталь', 'Коледино', 'Казань', 'Тула', 'Краснодар', 'Невинномысск', 'Подольск 3', 'Подольск')
entry_dest_warehouse_options.grid(row=2, column=1, padx=10, pady=2)

tk.Label(root, text="Тип поставки").grid(row=3, column=0, padx=10, pady=2)
delivery_type_var = tk.StringVar(root)
delivery_type_options = ttk.Combobox(root, textvariable=delivery_type_var)
delivery_type_options['values'] = ('Монопаллета', 'Короба')
delivery_type_options.grid(row=3, column=1, padx=10, pady=2)

tk.Label(root, text="Дата поставки (dd-mm-yyyy)").grid(row=4, column=0, padx=10, pady=2)
entry_delivery_date = DateEntry(root, date_pattern='dd.mm.yyyy')
entry_delivery_date.grid(row=4, column=1, padx=10, pady=2)

tk.Label(root, text="Количество паллет").grid(row=5, column=0, padx=10, pady=2)
entry_amount_var = tk.StringVar()
entry_amount = tk.Entry(root, textvariable=entry_amount_var, state='readonly')
entry_amount.grid(row=5, column=1, padx=10, pady=2)

top_barcode_path = None
top_barcode_label = tk.Label(root, text="Top Barcode not selected")
top_barcode_label.grid(row=6, column=0, columnspan=2, pady=2)

top_barcode_button = tk.Button(root, text="Top Barcode", command=browse_top_barcode)
top_barcode_button.grid(row=7, column=0, columnspan=2, pady=2)

pallet_fields = []

add_pallet_button = tk.Button(root, text="Добавить Паллет", command=add_pallet_field)
add_pallet_button.grid(row=8, columnspan=2, pady=10)

submit_button = tk.Button(root, text="Отправить", command=submit_form)
submit_button.grid(row=100, columnspan=2, pady=20)

if __name__ == '__main__':
    root.mainloop()
